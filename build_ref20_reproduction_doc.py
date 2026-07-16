from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from pathlib import Path

OUT = Path('参考文献20_丝蛋白音乐化_复现操作手册.docx')

BLUE = '2E74B5'; DARK = '1F4D78'; LIGHT = 'E8EEF5'; GREY = 'F2F4F7'

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); mar = tcPr.first_child_found_in('w:tcMar')
    if mar is None:
        mar = OxmlElement('w:tcMar'); tcPr.append(mar)
    for side, val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = mar.find(qn('w:'+side))
        if node is None: node = OxmlElement('w:'+side); mar.append(node)
        node.set(qn('w:w'), str(val)); node.set(qn('w:type'), 'dxa')

def set_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.find(qn('w:tcW'))
    if tcW is None: tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(dxa)); tcW.set(qn('w:type'),'dxa')

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c,LIGHT); set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.bold=True
        if widths: set_width(c,widths[i])
    for row in rows:
        cells=t.add_row().cells
        for i,x in enumerate(row):
            cells[i].text=x; set_cell_margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths: set_width(cells[i],widths[i])
    doc.add_paragraph()
    return t

def add_link_para(doc, label, url, note=''):
    p=doc.add_paragraph(style='Normal'); r=p.add_run(label+': '); r.bold=True
    h=OxmlElement('w:hyperlink'); rel=doc.part.relate_to(url,'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',is_external=True); h.set(qn('r:id'),rel)
    run=OxmlElement('w:r'); rp=OxmlElement('w:rPr'); col=OxmlElement('w:color'); col.set(qn('w:val'),'0563C1'); rp.append(col); u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rp.append(u); run.append(rp); txt=OxmlElement('w:t'); txt.text=url; run.append(txt); h.append(run); p._p.append(h)
    if note: p.add_run(' — '+note)

def bullet(doc, text):
    p=doc.add_paragraph(style='List Bullet'); p.add_run(text); return p

def step(doc, title, text):
    p=doc.add_paragraph(style='List Number'); p.add_run(title+'：').bold=True; p.add_run(text); return p

doc=Document()
sec=doc.sections[0]; sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.75); sec.left_margin=Inches(0.85); sec.right_margin=Inches(0.85); sec.header_distance=Inches(0.35); sec.footer_distance=Inches(0.35)
styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); normal.font.size=Pt(10.5); normal.paragraph_format.space_after=Pt(5); normal.paragraph_format.line_spacing=1.18
for name,size,color,before,after in [('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',11.5,DARK,8,4)]:
    s=styles[name]; s.font.name='Calibri'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=True; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
header=sec.header.paragraphs[0]; header.text='参考文献20复现操作手册'; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs: r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(100,100,100)
footer=sec.footer.paragraphs[0]; footer.text='仅数字复现无需湿实验；生物实验须在具备相应生物安全与实验资质的实验室开展。'; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
for r in footer.runs: r.font.size=Pt(8); r.font.color.rgb=RGBColor(100,100,100)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(38); p.paragraph_format.space_after=Pt(10)
r=p.add_run('参考文献20：丝蛋白-音乐转译\n复现操作手册'); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=RGBColor.from_string('0B2545'); r.font.name='Calibri'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('基于 Wong et al., Nano Today 7 (2012) 488-495').italic=True
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('版本 1.0 | 2026-07-15 | 面向数字声化复现与可选材料实验验证').font.color.rgb=RGBColor(90,90,90)
doc.add_paragraph()
table(doc,['先做什么','完成标准','预计投入'],[
    ('第 1 阶段：数字复现（推荐）','获得 4 个音频、按固定映射重建 MIDI/乐谱，并与原曲比较','1-2 天'),
    ('第 2 阶段：音乐分析','给出 A/B 区块、音域、时值和主题重复的可复核标注','1 天'),
    ('第 3 阶段：材料验证（可选）','在合作实验室验证 HAB3 与 HA3B 的成纤维差异','数周至数月')
],[2300,4800,1200])

doc.add_heading('1. 复现范围与成功判据',1)
doc.add_paragraph('本文献不是“每种氨基酸直接映射为一个音符”的通用算法论文，而是将丝蛋白区块、组装行为和音乐层级对应起来的案例。最稳妥的复现是先做到“补充材料的音乐重建与听觉/乐谱验证”，再决定是否进入蛋白表达、微流控纺丝和 DPD 模拟。')
table(doc,['层级','输入','交付物 / 判据'],[
    ('A. 资料复核','原文、补充 PDF、4 个 MP3','文件名、时长、曲目-构建体关系一致'),
    ('B. 声化重建','补充材料的音高键、时值常量、区块顺序','每首导出 MIDI、MusicXML、WAV；人工审听可辨认共同 His 前缀与 A/B 主题'),
    ('C. 定量比对','原始 MP3 与重建音频','音高序列、节奏位置、总时长的差异表；差异逐项标注为原文未规定的作曲选择'),
    ('D. 材料实验（可选）','完整构建体 DNA/蛋白、微流控条件','HAB3 有连续纤维，HA3B 以团簇/不连续组装为主；不是仅凭音乐判断')
],[1450,2500,4350])
doc.add_paragraph('关键结论：HAB3 为成纤维构建体；HA3B 为不成纤维构建体。补充材料与 PMC 版本对四首曲目的编号和名称给出更清晰的对应关系，复现记录应采用该对应关系。')

doc.add_heading('2. 数据、软件与目录准备',1)
doc.add_heading('2.1 必须下载的数据',2)
for x in ['主文 PDF：Materials by design: Merging proteins and music（DOI: 10.1016/j.nantod.2012.09.001）。','补充材料 PDF（含音高键、范围、常量和作曲规则）。','四个原始 MP3：01 Spinning Melody 1_ Vivo；02 Spinning Melody 2_ Andantino；03 Spinning Melody 3_ Vivo；04 Spinning Melody 4 Andantino; Bright。','用于完全材料复现实验的额外资料：参考文献[14]的完整构建体序列、表达/纯化和微流控条件。当前论文不能单独提供这些可订购的 DNA 序列与所有工艺参数。'] : bullet(doc,x)
doc.add_heading('2.2 软件（免费优先）',2)
table(doc,['用途','推荐工具','具体任务'],[
    ('乐谱与 MIDI','MuseScore 4','新建长笛单声部，输入映射后的音高/时值，导出 MusicXML、MIDI、PDF'),
    ('听觉/频谱检查','Audacity 或 Sonic Visualiser','波形、频谱、节拍标记、与原 MP3 的 A/B 段落对照'),
    ('可复现脚本','Python 3.11+；music21、mido、pretty_midi','将 CSV 音符事件确定性导出 MIDI；生成日志与校验表'),
    ('音频渲染','MuseScore 内置合成器或 FluidSynth + 长笛 SoundFont','所有曲目固定同一乐器、采样率和响度设置'),
    ('可选模拟','LAMMPS 或自建 DPD 工作流','仅在获得论文[15]的参数、边界条件与验证数据后实施')
],[1400,2600,4300])
doc.add_heading('2.3 建议目录结构',2)
for x in ['ref20_reproduction/00_sources/（原文、补充 PDF、下载日期说明）','ref20_reproduction/01_original_audio/（4 个原始 MP3，只读保存）','ref20_reproduction/02_annotations/（每首曲目的 CSV、段落标注、人工决策日志）','ref20_reproduction/03_scores/（.mscz、.musicxml、.mid、.pdf）','ref20_reproduction/04_rendered_audio/（重建 WAV/MP3）','ref20_reproduction/05_comparison/（截图、差异表、审听记录）'] : bullet(doc,x)

doc.add_heading('3. 已公布的映射规则：先固定不可变项',1)
doc.add_paragraph('补充材料给出的规则可直接作为复现规范。注意：它既包含确定性映射，也保留“为审美、流畅度和动态性选择高低音区”等人为作曲自由度。因此必须把每一次自由选择写入 decisions.csv，不能把它误称为论文给出的唯一答案。')
table(doc,['规则','实施方式','复核点'],[
    ('字母→音高键','M=E；H=B；S=B♭；G=G；L=A♭；V=D♭；P=E♭；R=D；K=F♯；E=高八度 E；F=F；Q=低八度 E♭；D=低八度 D；Y=高八度 F♯。','CSV 使用科学音高记法（如 B4）；注明选定八度。'),
    ('区块主题','A 为 12 音符、疏水/富 Ala 区块；B 为 20 音符、亲水区块。按构建体的 A/B 次数串联。','A 被实现为更仪式化/攻击性；B 更抒情、内部变化更多。'),
    ('共同前缀','四曲均先演奏同一段 48 音符 His-tag 乐句。','所有重建曲在事件 1-48 绝对一致。'),
    ('TS linker','C-B♭；两个音各为 6 拍长。','每个 linker 的 pitch 与 duration 完全一致。'),
    ('音域与移位','限制在长笛可演奏音域内；不同曲可整体/局部移位。','写入每段 transpose_semitones；不得超出音域。'),
    ('时值倾向','部分音符跨曲保持标志性时值。','未明确时值须标“推定”，并用原 MP3 对齐。')
],[1400,4200,2700])
doc.add_paragraph('建议的机器可读事件表字段：piece_id, event_index, source_symbol, block, pitch_name, midi_note, onset_beats, duration_beats, articulation, octave_choice, rule_source, confidence, note。')

doc.add_heading('4. 可执行的数字复现步骤（主路线）',1)
step(doc,'建立可追溯资料库','下载主文、补充 PDF 和 4 个 MP3；在 sources_manifest.csv 记录来源 URL、访问日期、文件名、文件大小、SHA-256。原始文件不做覆盖编辑。')
step(doc,'建立曲目-构建体对照表','固定为：01=HAB3 Melody 1（成纤维）；02=HA3B Melody 1（不成纤维）；03=HAB3 Melody 2（成纤维、较长）；04=HA3B Melody 2（不成纤维、较长）。')
step(doc,'配置软件','安装 MuseScore 4、Audacity/Sonic Visualiser 和 Python。新建一支单声部长笛工程，拍号/速度先留为“待校准”，不要凭感觉设定。')
step(doc,'转录共同前缀','以补充材料的 48 音符 His-tag 规则为基线；若补充 PDF 未列出完整时值，播放 01 并以小节线和波形峰值手工标注 onset/duration。保存 prefix_48.csv。')
step(doc,'构建 A 和 B 主题表','依补充材料的字母→音高键写出 A（12 个）和 B（20 个）音高序列，分别存为 motif_A.csv 与 motif_B.csv。每个音符保留原始氨基酸字母，不要只保存 MIDI 数值。')
step(doc,'按构建体拼接','生成 HAB3 = H + A + B + B + B；HA3B = H + A + A + A + B 的事件级草稿。按补充材料中 linker 的实际位置插入 TS；若位置未明，标为待从原曲/乐谱确认。')
step(doc,'从 MP3 校准自由参数','在 Sonic Visualiser 导入原 MP3 与草稿 MIDI；逐段校准八度、节奏、重音和速度。每次修改在 decisions.csv 说明证据（时间戳、听觉判断或补充材料条目）。')
step(doc,'生成乐谱与音频','将定稿 CSV 通过 music21/mido 生成 MIDI，再导入 MuseScore 排版为单声部长笛谱。以固定 SoundFont/音色渲染 WAV；再编码为 MP3。')
step(doc,'逐项验收','检查：48 音符前缀一致；TS 均为 C-B♭且各 6 拍；HAB3 的 B 主题次数多于 HA3B；四首的构建体标签没有颠倒；乐谱、MIDI、WAV 的版本号相同。')
step(doc,'输出复现包','提交 README、sources_manifest.csv、四首 .mscz/.musicxml/.mid/.wav、events.csv、decisions.csv、comparison_report.pdf。README 必须清楚区分“文献明确规定”与“根据音频推定”。')

doc.add_heading('5. 验证与报告模板',1)
table(doc,['检查项','通过标准','记录位置'],[
    ('文件完整性','4 个原始音频均可播放，哈希已记录','00_sources/sources_manifest.csv'),
    ('映射一致性','所有已公布字母映射无冲突；TS= C-B♭、各 6 拍','02_annotations/rule_audit.csv'),
    ('共同前缀','四曲事件 1-48 的 pitch/duration 一致','02_annotations/prefix_check.csv'),
    ('构建体结构','HAB3 与 HA3B 的 A/B 次数和标签正确','02_annotations/construct_check.csv'),
    ('听觉对齐','每首至少标注 5 个锚点，记录偏差和原因','05_comparison/alignment_notes.md'),
    ('可复跑性','删除输出后，运行脚本可重建 MIDI/WAV','README 的一条命令与环境锁定文件')
],[1600,4000,2700])
doc.add_paragraph('不要采用“音频相似度单一阈值”作为唯一成功标准：原作者明确保留了音域、流动性和动态性的作曲选择。优先验证离散规则、构建体关系和可审计的人工决策。')

doc.add_heading('6. 材料实验与模拟：范围、缺口与安全边界',1)
doc.add_paragraph('该部分不应被视为可由本文献单独完成的 SOP。原文将部分关键工作归到既有文献和“in preparation”的模型，因此缺少完整构建 DNA、浓度、pH、流量、芯片尺寸、DPD 势参数、随机种子和原始数据。以下为实施前清单，而非未经验证的实验配方。')
table(doc,['工作包','必须先获得','最小观察结果'],[
    ('蛋白构建与表达','HA3B/HAB3 的确切 DNA/氨基酸序列、载体、宿主、纯化方案及实验室审批','SDS-PAGE/质谱确认产物身份和纯度'),
    ('微流控纺丝','芯片几何、核心/鞘液配方、浓度、流量、收集与干燥条件','显微图像显示连续纤维或非连续团簇'),
    ('结构/力学表征','样本数、长度/直径测量法、拉伸速率、统计计划','HAB3 与 HA3B 的纤维形成差异及置信区间'),
    ('DPD 模拟','所有 bead 映射、相互作用、剪切流条件、边界、时间步和对照数据','可重现的团簇连通性/贯通网络指标')
],[1600,4300,2600])
doc.add_paragraph('建议决策门：未取得完整构建体序列和微流控参数前，停止湿实验；未取得 DPD 参数或可验证的原始轨迹前，停止将模拟结果称为“论文复现”。可以先完成数字复现，不受这些缺口影响。')

doc.add_heading('7. 可用资源、开源项目与成品页面',1)
doc.add_paragraph('本次检索未发现该 2012 论文作者公开发布、可直接运行的“丝蛋白→音乐”专用代码仓库或交互网站。其可直接复现的成品是 PMC 托管的补充音频与补充 PDF。下列资源按用途区分，避免把后续相关项目误当作本论文的原始实现。')
add_link_para(doc,'论文全文与成品音频/补充材料','https://pmc.ncbi.nlm.nih.gov/articles/PMC3752788/','最重要：页面 Associated Data 中有 4 个 MP3 与补充 PDF。')
add_link_para(doc,'论文书目信息（PubMed）','https://pubmed.ncbi.nlm.nih.gov/23997808/','核对 DOI、作者、期刊与图注。')
add_link_para(doc,'开放的音乐结构分析工具','https://www.sonicvisualiser.org/','Sonic Visualiser：音频可视化、标注、导入 MIDI；并非本文作者代码。')
add_link_para(doc,'开放的算法作曲工具','https://github.com/ideoforms/isobar','Python MIDI/OSC 模式生成库，MIT 许可；可用于将事件表导出为 MIDI，并非本文作者代码。')
add_link_para(doc,'后续相关开源项目','https://github.com/lamm-mit/AttentionCrossTranslation','2023 年跨域音乐-蛋白深度学习项目；与本文主题相关但不是 2012 规则的复现代码。')
doc.add_paragraph('使用提示：下载的补充音频用于研究、比较与教学时应保留原始署名、文章 DOI 和来源链接；不要将它们或重建版本暗示为作者官方新发布的程序。')

doc.add_heading('8. 参考依据与复现限制',1)
for x in ['Wong JY, McDonald J, Taylor-Pinney M, Spivak DI, Kaplan DL, Buehler MJ. Materials by design: Merging proteins and music. Nano Today. 2012;7(6):488-495. DOI: 10.1016/j.nantod.2012.09.001。','该文补充材料：含四首音频与“Translation from silk to music”规则；其中把字母映射、长笛音域、TS linker、共同 His 前缀列为关键常量。','Kinahan et al. Tunable Silk: Using Microfluidics to Fabricate Silk Fibers with Controllable Properties. Biomacromolecules. 2011;12(5):1504-1511。用于追溯微流控纺丝背景，但不能替代取得本研究确切构建体和参数。'] : doc.add_paragraph(x,style='List Bullet')
doc.add_paragraph('限制声明：本文将可检索到的补充规则转写为可执行流程。文献没有公布全部作曲决策和完整 DPD 参数，因此“完全逐音一致”和“完全实验/模拟重演”均需要向作者、补充档案或原始相关研究进一步索取资料。')

doc.save(OUT)
print(OUT.resolve())
