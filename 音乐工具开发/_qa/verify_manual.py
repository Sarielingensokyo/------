from pathlib import Path

from docx import Document


root = Path(__file__).resolve().parents[1]
path = root / "BioSound_GVR生物物理表现层增强版平台说明.docx"
document = Document(path)
texts = [paragraph.text for paragraph in document.paragraphs]
for table in document.tables:
    for row in table.rows:
        texts.extend(cell.text for cell in row.cells)
all_text = "\n".join(texts)

required = [
    "生物物理调式映射（推荐）",
    "文献氨基酸映射（复现）",
    "可逆十二音列载体（实验）",
    "Shrake–Rupley",
    "surface_wetness",
    "CC93",
    "裸音高",
    "NMA",
    "SoundFont 采样器、音色库与许可",
    "Clarinet-20190818.sf2",
    "TimGM.sf2",
    "44.1 kHz",
]
stale = ["摘要前 12 个字节", "平台默认使用十二音列"]
missing = [term for term in required if term not in all_text]
unexpected = [term for term in stale if term in all_text]
headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]

print(
    {
        "file": str(path),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "sections": len(document.sections),
        "headings": len(headings),
        "missing": missing,
        "stale": unexpected,
        "bytes": path.stat().st_size,
    }
)
assert not missing
assert not unexpected
assert len(headings) >= 20
