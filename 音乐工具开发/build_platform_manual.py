from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "BioSound_GVR可逆十二音列版平台功能与原理说明.docx"

# compact_reference_guide tokens + named CJK font override
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "102A43"
TEAL = "087F8C"
MUTED = "627D98"
PALE_BLUE = "E8EEF5"
PALE_TEAL = "E4F1EF"
PALE_GOLD = "FFF4D8"
PALE_RED = "FDECEC"
WHITE = "FFFFFF"
TABLE_WIDTH = 9360
TABLE_INDENT = 120
CJK_FONT = "Microsoft YaHei"


def set_font(run, size=None, color=None, bold=None, italic=None, ascii_font="Calibri"):
    run.font.name = ascii_font
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths):
    assert sum(widths) == TABLE_WIDTH
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def style_table_text(table, header=True, size=9.5):
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_font(run, size=size, color=NAVY, bold=(header and row_index == 0))
            if header and row_index == 0:
                shade(cell, PALE_BLUE)
    if header and table.rows:
        set_repeat_table_header(table.rows[0])


def add_table(doc, headers, rows, widths, size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = str(value)
    set_table_geometry(table, widths)
    style_table_text(table, header=True, size=size)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, label, text, fill=PALE_TEAL, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(f"{label}  ")
    set_font(r, size=10.5, color=color, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.5, color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), size=11, color=NAVY)
    return p


NUMBER_COUNTER = 0


def reset_numbers():
    global NUMBER_COUNTER
    NUMBER_COUNTER = 0


def add_number(doc, text):
    global NUMBER_COUNTER
    NUMBER_COUNTER += 1
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(f"{NUMBER_COUNTER}.  "), size=11, color=NAVY, bold=True)
    set_font(p.add_run(text), size=11, color=NAVY)
    return p


def add_para(doc, text, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, size=11, color=NAVY, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r, size=11, color=NAVY, italic=italic)
    else:
        set_font(p.add_run(text), size=11, color=NAVY, italic=italic)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F6F7")
    p_pr.append(shd)
    for line in text.splitlines():
        r = p.add_run(line)
        set_font(r, size=9.2, color=NAVY, ascii_font="Consolas")
        r.add_break()
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = paragraph.add_run("第 ")
    set_font(prefix, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    field_run = paragraph.add_run()
    set_font(field_run, size=9, color=MUTED)
    field_run._r.append(fld_begin)
    field_run._r.append(instr)
    field_run._r.append(fld_separate)
    result_run = paragraph.add_run("1")
    set_font(result_run, size=9, color=MUTED)
    end_run = paragraph.add_run()
    set_font(end_run, size=9, color=MUTED)
    end_run._r.append(fld_end)
    suffix = paragraph.add_run(" 页")
    set_font(suffix, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(hp.add_run("BioSound GVR  |  多声部管弦乐版功能与原理说明"), size=9, color=MUTED)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    set_font(p.add_run("BIOSOUND GVR"), size=11, color=TEAL, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run("多声部管弦乐版\n平台功能与原理说明"), size=27, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    set_font(p.add_run("从生物序列、三维结构与多组学指标到可追溯的六声部古典编配"), size=14, color=DARK_BLUE)
    add_callout(
        doc,
        "适用对象",
        "音乐初学者、生物信息学研究者、跨学科创作者与平台开发者。本文档依据《Transform from genes to music》综述的方法框架，覆盖全部界面功能、模态特异映射、六声部配器、GVR 约束、文件导出、科研边界和二次开发接口。",
        fill=PALE_TEAL,
    )
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("版本 2.0  |  2026 年 7 月 16 日  |  本地离线多声部版"), size=10.5, color=MUTED, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("工作目录：C:\\Users\\34591\\Desktop\\音乐基因转化\\音乐工具开发"), size=9, color=MUTED)
    doc.add_page_break()


def build():
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    doc.add_heading("阅读导航", level=1)
    add_para(doc, "本手册从“如何立即使用”开始，再逐步解释数据如何变成音乐。只想生成作品的读者可重点阅读第 1–4 章；需要科研复现的读者应继续阅读第 5–11 章；准备修改代码的开发者请阅读第 12–14 章与附录。")
    for item in [
        "第 1 章：平台定位、能力边界与快速启动",
        "第 2 章：网页界面与完整操作流程",
        "第 3 章：文本、FASTA、PDB 与 CSV 输入规范",
        "第 4 章：试听、审计与全部导出功能",
        "第 5 章：序列与理化特征提取原理",
        "第 6 章：PDB 空间声像、接触图与粗粒化 NMA",
        "第 7 章：表达、表观、代谢、质谱与关联数据声学化",
        "第 8 章：六声部古典编配、独立谱表与音频合成",
        "第 9 章：十二音列与 Generate–Verify–Repair",
        "第 10 章：溯源、可复现性与科研使用",
        "第 11 章：多类输入的完整示例工作流",
        "第 12–14 章：故障排查、软件架构、测试与扩展",
        "第 16 章：综述方法矩阵、已实现功能与尚未实现边界",
        "附录：氨基酸映射、Trace 字段、GVR 报告与发布清单",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "最重要的科学边界", "本平台输出的是规则明确、可审计的数据声学化模型，不是蛋白质或细胞直接发出的真实声音，也不替代结构测定、统计检验或临床判断。", fill=PALE_GOLD)

    doc.add_heading("1. 平台定位与快速启动", level=1)
    doc.add_heading("1.1 平台做什么", level=2)
    add_para(doc, "BioSound GVR 把一维生物序列、PDB 三维坐标、质谱峰或多组学指标编译为可同时发声的古典室内乐事件。默认编制包括生物主旋律、弦乐对位、大提琴结构低音、圆号和声场、中提琴内声部与竖琴结构重音。每个派生事件都保存主事件编号、来源位置、输入符号、理化特征、预期音级、音乐功能和 GVR 状态。")
    add_para(doc, "平台的核心价值不是把字母机械替换成音符，也不是在一条旋律上不断切换乐器，而是采用综述提出的层级框架：最重要的生物组织轴控制最重要的音乐组织轴；序列/坐标决定时间，身份或数值决定前景音高，结构与聚类决定声部和形式，质量与不确定性决定滤波、力度和清晰度，验证器阻止错误映射进入最终文件。")
    doc.add_heading("1.2 一分钟启动", level=2)
    reset_numbers()
    for step in [
        "打开工作目录，双击“启动平台.bat”。",
        "第一次启动若尚无本地环境，程序会自动创建 .venv 并安装依赖；以后启动不会重复询问。",
        "浏览器自动打开 http://localhost:8501。平台不要求邮箱、注册、登录或云端账户。",
        "从左侧选择示例、粘贴序列或上传文件，设置音高策略和速度，然后点击“生成并验证”。",
        "在主页面试听，并按需下载 WAV、MIDI、MusicXML、PDF、Trace CSV 或 GVR JSON。",
    ]:
        add_number(doc, step)
    add_callout(doc, "隐私", "Streamlit 以 headless 本地模式运行，已关闭浏览器统计收集。上传或粘贴的数据只在当前计算机内处理。", fill=PALE_TEAL)
    doc.add_heading("1.3 能力边界", level=2)
    add_bullet(doc, "可以：教学演示、规则比较、数据质控的听觉辅助、跨学科创作、可复现原型研究。")
    add_bullet(doc, "不应直接用于：疾病诊断、结构预测结论、细胞质量的唯一判断、把悦耳程度解释为生物功能。")
    add_bullet(doc, "网页试听是可叠加的轻量古典乐器频谱近似；出版级音质应把多轨 MIDI 或多谱表 MusicXML 送入 Muse Sounds、专业管弦乐采样库或真实演奏。")

    doc.add_heading("2. 网页界面与完整操作流程", level=1)
    doc.add_heading("2.1 左侧输入与作曲设置", level=2)
    add_table(doc, ["控件", "作用", "建议"], [
        ("数据来源", "示例数据、粘贴序列、上传文件三选一", "初次使用先选示例"),
        ("序列名称", "为纯文本序列生成 FASTA 标题", "使用样本号或基因/蛋白名称"),
        ("序列类型", "自动识别、蛋白质、DNA、RNA", "含模糊字符时手动指定"),
        ("记录/链", "多 FASTA 记录或多 PDB 链的选择器", "每次处理一个记录"),
        ("音高策略", "生物物理映射、文献氨基酸映射、可逆十二音列编解码", "科研比较时保存设置"),
        ("调式", "多利亚、五声、自然小调、半音阶", "文献映射模式主要影响非蛋白输入"),
        ("十二音列形式", "P、I、R、RI", "仅十二音列模式启用"),
        ("速度/拍号", "控制播放时间与小节组织", "同批比较保持一致"),
        ("最多生成事件", "对非可逆模式的超长输入等间隔抽样", "可逆模式自动禁用抽样"),
        ("古典编配层数", "从 1 到 6 逐步加入对位、低音、圆号、中提琴和竖琴", "需要完整效果时选 6"),
        ("对位独立度", "控制小提琴延迟、反向音区与主旋律的差异", "初学者建议 0.70"),
        ("随机种子", "固定试听合成细节；不参与可逆音列编码", "同一实验固定种子"),
        ("PDB NMA", "有 CA 坐标时计算相对简正模式", "大结构会自动抽样"),
    ], [1750, 4810, 2800], size=9.2)
    doc.add_heading("2.2 生成以后看到什么", level=2)
    add_para(doc, "页面顶部给出输入点数、总音乐事件数、GVR 是否通过、自动修复次数、音频时长和独立声部数。下方四个标签页分别承担“听和下载”“逐音符追踪”“规则验证”“科学解释”四种任务。")
    add_bullet(doc, "试听与导出：播放可同时发声的立体声 WAV、查看当前编制与每声部事件数、查看特征曲线并下载全部产物。")
    add_bullet(doc, "生物—音乐轨迹：逐行查看声部、音乐功能、主事件 ID、来源位置、符号、起拍、时值、MIDI、音级、声像、乐器、状态和映射理由。")
    add_bullet(doc, "GVR 与映射：检查硬规则是否通过、查看修复日志、核对氨基酸音高配置。")
    add_bullet(doc, "科学边界：查看 NMA 参数、CSV QC 假设和论文方法在本平台中的实现边界。")
    add_para(doc, "页面另设“从平台产物还原 DNA / RNA / 蛋白质序列”面板，可严格读取平台生成的 GVR JSON、MusicXML 或 MIDI。WAV 只用于试听，不作为无损解码载体。")

    doc.add_heading("3. 输入格式与解析规则", level=1)
    doc.add_heading("3.1 直接粘贴序列", level=2)
    add_para(doc, "在“数据来源”选择“粘贴序列”。可直接粘贴一行或多行纯字母，也可以粘贴完整 FASTA。纯文本会自动加上由“序列名称”生成的 FASTA 标题；空格、数字和常见换行会在解析时清除。")
    add_code(doc, ">my_protein\nMKTIIALSYIFCLVFADYKDDDDK\n\n# 纯序列也可以：\nACGTACGTACGT")
    add_para(doc, "自动识别逻辑：字符集合仅属于 A/C/G/T/N 时判为 DNA；仅属于 A/C/G/U/N 时判为 RNA；其他情况判为蛋白质。若蛋白质只由 A、C、G、T 等少数氨基酸组成，自动识别可能误判，必须手动选择“蛋白质”。")
    doc.add_heading("3.2 FASTA", level=2)
    add_para(doc, "支持 .fa、.fasta、.faa、.fna。多记录 FASTA 会在“记录/链”控件中列出；平台每次生成其中一个记录。标题作为记录名，字符位置从 1 开始写入 Trace CSV。")
    doc.add_heading("3.3 PDB", level=2)
    add_para(doc, "PDB 解析器读取 ATOM/HETATM 中的 CA 原子，忽略非 A 构象的 alternate location，并按链和残基号去重。三字母残基转换为一字母代码；不能识别的残基记为 X。HELIX 与 SHEET 记录用于二级结构类别，其余残基记为 coil。")
    add_callout(doc, "注意", "没有 CA 原子、坐标列格式不合法或仅含配体的 PDB 无法进入蛋白质声学化。mmCIF 当前版本尚未支持。", fill=PALE_GOLD)
    doc.add_heading("3.4 CSV 的六种解释路径", level=2)
    reset_numbers()
    add_number(doc, "若存在 sequence、seq、protein_sequence 或 dna_sequence 列：把每行解析为独立序列。")
    add_number(doc, "若同时存在 mz/m-z 与 intensity/abundance 类列：按质谱峰解释。")
    add_number(doc, "若存在 p/pvalue、染色体或坐标列：按 GWAS/EWAS 类关联景观解释，位置进入时间，-log10(p) 进入显著性。")
    add_number(doc, "若存在 methyl、CpG、H3K、ChIP 或 chromatin 类数值列：按表观调控轨迹解释。")
    add_number(doc, "若存在 metabolite/compound 与 abundance/intensity：按代谢丰度表解释，并在有置信度列时保留注释不确定性。")
    add_number(doc, "以上均不匹配时，把数值矩阵启发式解释为转录组表达矩阵，并假定行为细胞、列为基因。")
    add_para(doc, "CSV 可用 UTF-8 或 GB18030。缺失值与无穷值会被替换为可计算数值。自动分类依赖列名，只是入口级启发式；正式分析前必须核对模态、矩阵方向、标准化、字段含义和单位。尤其是单细胞表达数据常以基因为行、细胞为列，需要先转置。")

    doc.add_heading("4. 试听、审计与导出功能", level=1)
    doc.add_heading("4.1 WAV 试听", level=2)
    add_para(doc, "网页生成 22.05 kHz、16 位、双声道 WAV。六个声部按各自起拍叠加，具有固定的古典乐器角色和基础舞台位置；PDB x 坐标在该基础上施加局部空间偏移。音频最长默认 150 秒，超出时仅试听被截断，MIDI、MusicXML 和 Trace 仍保留全部事件。")
    doc.add_heading("4.2 六类下载文件", level=2)
    add_table(doc, ["格式", "包含内容", "典型用途"], [
        ("WAV", "已渲染的双声道古典配器预览", "快速试听、演示"),
        ("MIDI", "格式 1：速度轨 + 每声部独立轨道、通道、乐器和声像", "DAW、专业音源、分轨编曲"),
        ("MusicXML", "每声部独立 Part、谱表、拍号、速度、歌词符号与乐器定义", "MuseScore/Sibelius 总谱与分谱"),
        ("PDF", "由本机 MuseScore 从 MusicXML 排版", "打印、分享、课堂展示"),
        ("Trace CSV", "逐音符来源与全部映射证书", "复现、审计、统计"),
        ("GVR JSON", "元数据、检查、修复、NMA、音频参数与完整载体音列", "研究归档、严格解码"),
    ], [1350, 4450, 3560], size=9.3)
    add_para(doc, "PDF 按钮需要本机安装 MuseScore 4。若检测不到 MuseScore，平台仍提供 MusicXML；用户可在任意兼容软件中手动导出 PDF。")

    doc.add_page_break()
    doc.add_heading("5. 序列与理化特征提取原理", level=1)
    doc.add_heading("5.1 蛋白质特征", level=2)
    add_para(doc, "每个氨基酸按 Kyte–Doolittle 疏水性表赋值 h，平台使用 h_norm = clip((h + 4.5) / 9, 0, 1) 归一化。酸性残基 D/E 的电荷取 -1，碱性 K/R 取 +1，H 取 +0.25；分子质量按常见残基质量表归一化到 75–205 Da。")
    add_table(doc, ["特征", "进入音乐的主要通道", "解释"], [
        ("氨基酸身份", "文献音高或调式级数", "保留字母/残基差异"),
        ("疏水性", "主旋律时值、内声部音区与和声色彩", "局部理化差异的听觉代理"),
        ("电荷", "主旋律八度、力度与十二音列目标音区", "正负电性的听觉对比"),
        ("残基质量", "当前写入特征层，供扩展", "可用于后续密度或力度实验"),
    ], [1900, 3350, 4110])
    doc.add_heading("5.2 DNA 与 RNA 特征", level=2)
    add_para(doc, "碱基身份 A/C/G/T(U)/N 被映射到所选调式的不同级数；G 或 C 的 GC 指示值控制高低音区，A 或 G 的 purine 指示值与 GC 加权形成 value = 0.65×GC + 0.35×purine。该设计让碱基身份、GC 与嘌呤/嘧啶差异进入不同音乐维度。")
    doc.add_heading("5.3 时值", level=2)
    add_para(doc, "PDB 有二级结构时：helix = 1 个四分音符拍，sheet = 1.5 拍，coil = 0.5 拍。只有序列时，疏水性高的残基通常为 0.5 拍，亲水性强的残基为 1 拍，中间区域为 0.75 拍。组学行的 detected_features 控制 0.5–1.5 拍。所有时值量化到 0.25 拍。")
    doc.add_heading("5.4 超长序列抽样", level=2)
    add_para(doc, "非可逆模式下，若输入长度 N 大于“最多生成事件”M，步长 stride = ceil(N/M)，按 0, stride, 2×stride…选择数据点。Trace CSV 保存真实 source_index。可逆十二音列模式绝不抽样：DNA/RNA 每块固定生成 12 个载体音符，蛋白质每 6 个残基生成 12 个载体音符；只有 WAV 试听可能按时长截断。")

    doc.add_heading("6. PDB 空间声像、接触图与粗粒化 NMA", level=1)
    doc.add_heading("6.1 三维坐标到左右声像", level=2)
    add_para(doc, "对每条链的 CA 坐标先减去质心。x 坐标按该链的 x 轴跨度归一化到 [-1, 1]：-1 为左，0 为中央，+1 为右。网页 WAV 使用等功率声像：left = cos((pan+1)π/4)，right = sin((pan+1)π/4)。因此耳机中左右能量连续变化，而不是简单开关。")
    doc.add_heading("6.2 接触度", level=2)
    add_para(doc, "任意两 CA 原子距离小于 8 Å 且不为自身时记为一次接触。每个残基的接触数量在链内归一化到 0–1，主要影响音区与大提琴配器：高接触度更容易进入较厚重、较集中声部。8 Å 是粗粒化接触图常见量级，本平台把它作为可配置 MVP 规则，不等同于化学键。")
    doc.add_heading("6.3 粗粒化 NMA", level=2)
    add_para(doc, "平台以 CA 节点建立单位弹簧各向异性网络模型。距离不超过默认 12 Å 的节点对建立弹簧，方向向量 d 形成 3×3 子块 -ddᵀ/||d||²，再装配 3N×3N Hessian。对 Hessian 求特征值，去除接近零的刚体模态，对正特征值 λ 取 sqrt(λ) 作为相对频率尺度。")
    add_para(doc, "为避免大蛋白造成内存和时间爆炸，最多抽取约 220 个残基；平台记录 sample_stride、sampled_residues 和 contacts。最低最多 24 个非零模式按对数比例映射到 55–1760 Hz，并以低音弦乐/圆号式频谱作为安静背景。")
    add_callout(doc, "不能过度解读", "单位弹簧模型没有真实力常数、原子质量、溶剂和温度标定；映射后的 Hz 只保持模式间相对对数比例，不是蛋白质绝对振动频率。", fill=PALE_RED, color="8A1C1C")

    doc.add_heading("7. 多组学、QC、质谱与统计景观声学化", level=1)
    doc.add_heading("7.1 表达矩阵 QC", level=2)
    add_para(doc, "对每一行计算 total_counts、detected_features 和 mitochondrial_fraction。线粒体列通过列名 MT- 或 MT_ 识别；线粒体比例 = 线粒体列总和 / 全部数值列总和。列方差最高的最多 20 个特征被当作 HVG 代理，行级 hvg_score 是这些列的均值再归一化。")
    add_para(doc, "线粒体比例越高，WAV 的低通截止频率越低：cutoff = 700 + (1 - mean_mito)×9000 Hz。HVG 分数提高前景主旋律的力度和清晰度；基因/细胞簇的复调潜力在当前版本以对位、低音与和声层派生表示，而不是让数千个基因同时无筛选地发声。")
    add_callout(doc, "QC 只是听觉审计", "阈值、基因命名和矩阵方向与实验平台密切相关。MVP 规则不能替代 Scanpy/Seurat 等成熟流程；适合先听出异常，再回到统计图和原始计数验证。", fill=PALE_GOLD)
    doc.add_heading("7.2 质谱", level=2)
    add_para(doc, "m/z 四舍五入后 mod 12 得到音级；归一化 m/z 决定音区，强度决定力度，竖琴承担前景峰序列与显著结构重音。该映射保留相对峰序列与强弱，但不能让听者直接反演精确 m/z；精确数值应从 Trace CSV 读取。")
    doc.add_heading("7.3 表观组学", level=2)
    add_para(doc, "基因组坐标沿音乐时间展开，甲基化或染色质信号控制主特征、力度、和声密度与结构层。综述强调不同组蛋白标记应成为不同声部或音色类别；当前 MVP 先实现单轨数值入口和六声部结构派生，多标记文件仍应先整理为可区分列，并在 Trace 中保留列名。低信号可形成稀疏区，高信号形成较密集、较突出的事件。")
    doc.add_heading("7.4 代谢组", level=2)
    add_para(doc, "代谢物丰度进入局部强度与音高，注释置信度进入 uncertainty，从而降低力度或清晰度。必须牢记丰度不等于通量：堆积可能来自产生增加、消耗降低或运输变化。当前版本尚未读取反应图与同位素示踪方向，因此生成的是代谢状态声景，而不是通量音乐。")
    doc.add_heading("7.5 GWAS/EWAS 类关联数据", level=2)
    add_para(doc, "染色体/坐标决定事件时间，-log10(p) 决定显著性与前景突出程度，效应方向进入音区趋势。显著事件可触发竖琴结构标记，但强音、尖锐或不协和只表示统计显著或偏离参考，不代表致病性、因果性或机制已被证明。LD 区块、可信集和多基因分数的完整区域建模属于后续扩展。")

    doc.add_page_break()
    doc.add_heading("8. 音乐映射与古典配器", level=1)
    doc.add_heading("8.1 三种音高策略", level=2)
    add_table(doc, ["模式", "音级来源", "生物特征仍控制"], [
        ("文献氨基酸映射", "pitch_mapping.csv；保留 Spinning Melodies 已知映射", "电荷音区、时值、力度、配器、声像"),
        ("生物物理映射", "归一化值进入五声/多利亚/小调/半音阶", "结构、QC、接触度与空间位置"),
        ("可逆十二音列编解码", "每个序列块经进制转换与康托逆排名得到独立音列", "音区、节奏、力度、配器、声像"),
    ], [2200, 3680, 3480])
    add_para(doc, "文献表未覆盖 C、I、N、W，因为原始丝蛋白构建体未使用全部 20 种氨基酸。本平台以 extended_inference 明确标记扩展值，避免把软件推断误称为文献事实。")
    doc.add_heading("8.2 六声部古典编配", level=2)
    add_table(doc, ["声部", "固定乐器", "音乐功能", "生物来源"], [
        ("V1_melody", "长笛/双簧管/单簧管/竖琴等", "前景生物主旋律", "每个被抽样的原始序列或数据点"),
        ("V2_counterpoint", "小提琴", "延迟轮唱与反向音区对位", "主旋律事件的可追溯派生"),
        ("V3_bass", "大提琴", "分组边界与持续结构低音", "每个短语/数据窗口的锚点"),
        ("V4_horn_harmony", "圆号", "长时值结构和声场", "结构、功能或数值区块"),
        ("V5_viola_harmony", "中提琴", "和声三度与内声部连续性", "与圆号共享的区块锚点"),
        ("V6_harp_accents", "管弦竖琴", "结构边界、接触峰、显著性重音", "高接触、高值或短语起点"),
    ], [1900, 2200, 3000, 2260], size=8.9)
    add_para(doc, "“古典编配层数”不是音量旋钮，而是结构开关：1 只保留主旋律，2 加入小提琴，3 加大提琴，4 加圆号，5 加中提琴，6 加竖琴。这样既满足感知节制，也允许用户逐层听清每类映射。不同声部可跨声部同时发声，但同一声部内部仍按时间线排列。")
    doc.add_heading("8.3 九种古典乐器与持久通道", level=2)
    add_table(doc, ["内部名", "中文", "GM 程序（1-based）", "当前用途"], [
        ("flute", "长笛", "74", "DNA 主旋律"), ("clarinet", "单簧管", "72", "RNA/表达矩阵主旋律"),
        ("oboe", "双簧管", "69", "蛋白质主旋律"), ("bassoon", "大管", "71", "扩展低音配器接口"),
        ("french_horn", "圆号", "61", "关联数据主旋律或结构和声"), ("violin", "小提琴", "41", "派生对位"),
        ("viola", "中提琴", "42", "表观主旋律或内声部"), ("cello", "大提琴", "43", "结构低音"),
        ("orchestral_harp", "竖琴", "47", "质谱/代谢主旋律与结构重音"),
    ], [2100, 1450, 2050, 3760], size=9.1)
    add_para(doc, "每个 MusicXML Part 和 MIDI Track 只在开头设置一次持续乐器，不再在单通道上每个音符切换 Program Change。这是从旧版“单线换音色”到真实多声部总谱的关键改变。")
    doc.add_heading("8.4 对位、低音、和声与结构重音算法", level=2)
    add_bullet(doc, "对位：按 2–3 个主旋律事件抽取锚点，延迟 0.5 拍；普通调式采用三度及反向音区，十二音列采用 I 形式的对应位置。")
    add_bullet(doc, "低音：按约两小节形成窗口，主音/属音在大提琴音域持续到下一窗口。")
    add_bullet(doc, "圆号与中提琴：按更长区块形成持续和声场，分别承担根音/属功能与三度色彩。和声是艺术组织层，不应解释为分子稳定性。")
    add_bullet(doc, "竖琴：在二级结构变化、接触度峰、数值显著峰或短语起点演奏三音短琶音，作为可听的结构标记。")
    doc.add_heading("8.5 网页音色模型与厅堂反射", level=2)
    add_para(doc, "WAV 试听不依赖大型 SoundFont，而是用古典乐器典型谐波结构作轻量加法合成：长笛以基频为主，单簧管突出奇次谐波，双簧管具有较丰富高次谐波，弦乐采用递减谐波并加入轻微揉弦，圆号与大管突出低次谐波，竖琴使用指数衰减拨弦包络。六声部使用不同增益，混合后加入 37 ms 与 71 ms 的克制早期反射。")
    add_para(doc, "这能提供古典配器方向的本地预览，但不等于真实乐团采样。多轨 MIDI 可驱动 Muse Sounds、Kontakt、BBC Symphony Orchestra 等音源；多谱表 MusicXML 可在 MuseScore/Sibelius 中排总谱、分谱和进一步人工配器。")
    doc.add_heading("8.6 力度、声像与呼吸", level=2)
    add_para(doc, "力度限制在 MIDI 1–127 的安全范围内，各声部再按前景、对位、低音、和声和重音设置独立混合增益。默认每 24 个主旋律事件插入 0.25 拍呼吸空隙。大提琴和圆号偏左，中提琴偏右，小提琴偏右，竖琴接近中央；有 PDB 坐标时，x 轴声像以加权偏移叠加到舞台位置。")

    doc.add_heading("9. 十二音列与 GVR 原理", level=1)
    doc.add_heading("9.1 从生物序列得到十二音列", level=2)
    add_para(doc, "载体层不使用哈希映射。DNA 使用 A=0、C=1、G=2、T=3，RNA 把 U 置于第 4 位；每 12 个碱基构成一个四进制整数 N。蛋白质使用字母表 ACDEFGHIKLMNPQRSTVWY* 的 21 进制编号，每 6 个残基构成一个整数 N。4^12 = 16,777,216 < 12!，21^6 = 85,766,121 < 12!，因此每个合法块都能无碰撞地进入十二音列排列空间。")
    add_para(doc, "平台对 N 执行零起点词典序康托逆排名（Lehmer/factoradic unranking），得到 0–11 的唯一全排列。尾块不足时用 A 补齐，并保存 original_length、pad_length、字母表、块长、载体声部、行形式与版本。SHA-256 只作为解码后的完整性校验和，不承担映射。")
    add_para(doc, "P/I/R/RI 只改变呈现，不丢弃绝对音级：P 保持原排列，I 对每个音级取 -x mod 12，R 反转，RI 同时倒影和逆行。解码器先用记录的形式撤销变形，再计算康托排名。不能把第一音归零，否则会丢失载体转调信息。")
    add_code(doc, "P  = (p0, p1, ... p11)\nI  = (-p0, -p1, ... -p11) mod 12\nR  = reverse(P)\nRI = reverse(I)")
    doc.add_heading("9.2 事件证书", level=2)
    add_para(doc, "每个 MusicEvent 至少包含 event_id、voice_id、role、parent_event_id、source_index、source_label、symbol、onset、duration、midi、velocity、pan、timbre、expected_pc、mapping_rule、row_position、row_form、codec_block、is_codec_carrier、status 和 feature 字典。只有 V1_melody 且 is_codec_carrier=true 的事件进入解码；五个派生声部只负责艺术表现。")
    doc.add_heading("9.3 硬约束", level=2)
    add_table(doc, ["规则", "验证内容", "失败后处理"], [
        ("H_mapping", "实际音级 = 映射证书 expected_pc", "保持音级投影到合法 MIDI"),
        ("H_register", "每个声部在自己的可演奏音域与全局 36–96 内", "换八度，不改变音级"),
        ("H_timeline", "同一 voice_id 内不意外重叠；跨声部重叠合法", "只顺延发生冲突的单一声部"),
        ("H_duration", "时值 > 0", "设为 0.25 拍"),
        ("H_trace", "来源索引与标签存在", "不可安全修复则应拒绝"),
        ("H_row", "每个载体块与编码证书的音级顺序一致", "最终复检；失败不发布"),
        ("H_permutation", "每块恰好包含 0–11 且各出现一次", "严格拒绝，不静默猜测"),
        ("H_codec_domain", "DNA/RNA 排名 < 4^12；蛋白质排名 < 21^6", "严格拒绝；不截到同一最大值"),
        ("H_event_id", "所有事件编号全局唯一", "不可安全修复则拒绝发布"),
        ("H_parent", "每个派生声部事件引用有效主事件", "不可安全修复则拒绝发布"),
    ], [1800, 4660, 2900], size=9.2)
    doc.add_heading("9.4 Generate–Verify–Repair–Release", level=2)
    reset_numbers()
    add_number(doc, "Generate：规则生成器依据数据和设置提出事件。MVP 不需要在线 LLM。")
    add_number(doc, "Verify：对实际事件执行确定性硬规则，而不是读取生成器的自我声明。")
    add_number(doc, "Repair：按声部分组，只做保持生物约束的局部操作，如换八度保留音级、修正时值、顺延声部内部重叠。跨声部同时发声不会被错误消除。")
    add_number(doc, "Re-verify：对修复后的最终事件重新扫描。")
    add_number(doc, "Release：有任何剩余硬违规时不发布 WAV/MIDI/MusicXML。")
    add_callout(doc, "与论文一致的谨慎表述", "事件级约束通过不等于作品在所有音乐学层面都完全合法。平台报告的是已实现规则的通过状态，并保留可核查轨迹。", fill=PALE_TEAL)
    add_callout(doc, "严格模式", "未经修改且保留元数据的 JSON、MusicXML、MIDI 才承诺符号级精确往返。非法编辑默认报错；任何自动修复都会产生新候选序列，不能再称为恢复原始序列。", fill=PALE_GOLD)

    doc.add_heading("10. 溯源、可复现性与科研使用", level=1)
    doc.add_heading("10.1 最小复现包", level=2)
    add_para(doc, "一次正式实验至少保存：原始输入、pitch_mapping.csv、Trace CSV、GVR JSON、MusicXML/MIDI、软件版本、音高模式、调式、速度、拍号、事件上限、编配层数、对位独立度、随机种子、NMA 开关和截止距离。WAV 可重新渲染，因此不是唯一必需的源文件。")
    doc.add_heading("10.2 如何比较两个样本", level=2)
    add_bullet(doc, "固定音高模式、调式、速度、拍号、事件上限、种子和音域。")
    add_bullet(doc, "记录是否发生抽样，优先比较相同长度或采用一致窗口。")
    add_bullet(doc, "先比较 Trace 中的特征和事件统计，再讨论主观听感。")
    add_bullet(doc, "盲听时不要告诉听者样本标签，并收集可重复的评分维度。")
    doc.add_heading("10.3 可发表的准确措辞", level=2)
    add_para(doc, "推荐：本工具实施可配置、可追溯的规则声学化；对未修改且保留编解码元数据的符号载体，DNA/RNA 或蛋白质字符串可精确往返，并由确定性验证器检查事件级映射一致性。")
    add_para(doc, "不推荐：本工具无损还原蛋白质真实声音，或声音直接证明某种结构/疾病状态。")

    doc.add_heading("11. 多类输入的完整示例工作流", level=1)
    doc.add_heading("11.1 直接粘贴蛋白质", level=2)
    reset_numbers()
    for step in ["选择“粘贴序列”，输入名称。", "粘贴一字母蛋白质序列并手动选择“蛋白质”。", "音高策略选“文献氨基酸映射”，速度 96，拍号 4/4，编配层数 6，对位独立度 0.70。", "生成后先在“当前编制”确认 6 个独立声部，再在 Trace 核对 C/I/N/W 的扩展映射。", "下载多谱表 MusicXML、多轨 MIDI 和 Trace CSV，再用 MuseScore 生成 PDF。"]:
        add_number(doc, step)
    doc.add_heading("11.2 DNA/RNA", level=2)
    add_para(doc, "使用生物物理映射和多利亚/五声音阶。观察 GC 碱基是否进入更高音区；比较样本时固定根音、速度与事件上限。若序列只含 A/C/G/T 但实际上是蛋白质，必须手动指定类型。")
    doc.add_heading("11.3 PDB", level=2)
    add_para(doc, "上传 PDB，选择链，启用 NMA。戴耳机观察声像是否随 x 坐标移动；在科学边界标签查看 sample_stride、contacts、相对本征值与映射频率。正式报告中注明 8 Å 接触阈值和 12 Å NMA 弹簧截断。")
    doc.add_heading("11.4 单细胞表达 CSV", level=2)
    add_para(doc, "确保行为细胞、列为基因，线粒体基因以 MT- 或 MT_ 开头。生成后核对 mean_mitochondrial_fraction 和 low-pass cutoff。听到沉闷仅表示当前 QC 规则被触发，必须回到线粒体比例、小提琴图和过滤阈值复核。")
    doc.add_heading("11.5 十二音列", level=2)
    add_para(doc, "选择“可逆十二音列编解码”和 P/I/R/RI。生成后核对 H_row、H_permutation、H_codec_domain 与载体块数，并同时保存 GVR JSON。把 JSON、MusicXML 或 MIDI 上传到页面解码面板，应得到与输入完全一致的字符串。蛋白质模式恢复的是氨基酸序列，不能凭蛋白质反推出原始同义密码子。")
    doc.add_heading("11.6 表观、代谢或关联 CSV", level=2)
    add_para(doc, "先用清晰列名表达数据语义：表观信号使用 methyl/CpG/H3K/ChIP 类列名；代谢表至少包含 metabolite/compound 与 abundance；关联数据至少包含 pvalue 与 chromosome/position。生成后先核对 record.data_type 和输入元数据，再讨论声学结果。若自动分类错误，应整理列名或在导入前转换数据，而不是依据听感猜测模态。")

    doc.add_page_break()
    doc.add_heading("12. 故障排查", level=1)
    add_table(doc, ["现象", "常见原因", "处理"], [
        ("启动后没有网页", "8501 端口占用或环境未安装", "关闭旧实例；查看命令窗口；必要时重建 .venv"),
        ("第一次安装很慢", "需要下载 Streamlit 依赖", "保持联网；安装完成后可离线使用"),
        ("序列被识别成 DNA", "蛋白质字符集合碰巧只含 A/C/G/T", "在序列类型中手动选蛋白质"),
        ("PDB 报无 CA", "文件只有配体、mmCIF 或列格式异常", "换标准 PDB；确认存在 ATOM ... CA"),
        ("CSV QC 不合理", "行列方向或基因名不符", "转置矩阵；统一 MT- 命名；核对数值列"),
        ("PDF 按钮不可用", "没有 MuseScore 4 或路径不同", "下载 MusicXML 手动导出；安装/定位 MuseScore"),
        ("试听被截断", "总时长超过 150 秒", "提高速度、减少事件；完整谱仍在导出文件"),
        ("声音不够真实", "网页使用轻量音色模型", "把 MIDI/MusicXML 导入专业古典管弦乐音源"),
        ("仍像单声部", "编配层数为 1，或播放器忽略多轨/多谱表", "将层数设为 6；在 MuseScore/DAW 检查 6 个 Part/Track"),
        ("GVR 拒绝发布", "修复后仍有硬违规", "查看 violations_after 与 Trace，不要绕过门控"),
    ], [2300, 3330, 3730], size=9.0)
    add_callout(doc, "停止平台", "关闭启动时保留的命令窗口，或在窗口按 Ctrl+C。浏览器页面关闭并不会自动停止本地服务。", fill=PALE_GOLD)

    doc.add_heading("13. 软件架构与文件职责", level=1)
    add_para(doc, "平台采用数据模型—控制引擎—界面/输出的分层结构。解析器只负责把外部文件变成 BioRecord；特征层补充数值；映射器提出 MusicEvent；GVR 决定能否发布；合成与导出层只消费通过检查的事件。")
    add_code(doc, "Input → parsers.py → BioRecord → features.py\n      → mapping.py → MusicEvent[] → gvr.py\n      → synth.py / exporters.py → WAV / MIDI / MusicXML / PDF / Trace")
    add_table(doc, ["文件", "职责"], [
        ("app.py", "Streamlit 中文界面、文本输入、试听、下载和科学说明"),
        ("models.py", "BioRecord、MusicEvent、Violation、GVRReport 数据结构"),
        ("parsers.py", "FASTA、PDB、CSV 解析与格式识别"),
        ("features.py", "理化特征、PDB 声像/接触度、粗粒化 NMA"),
        ("codec.py", "DNA/RNA 四进制、蛋白质 21 进制、康托排名/逆排名与产物解码"),
        ("mapping.py", "模态特异音高、六声部派生、音域、十二音列与事件生成"),
        ("gvr.py", "按声部硬约束验证、局部修复、软齐奏提示和最终门控"),
        ("synth.py", "古典乐器频谱近似、多声部混合、舞台声像、QC 滤波、WAV"),
        ("exporters.py", "多轨 MIDI、多 Part MusicXML、MuseScore PDF、GVR JSON"),
        ("pipeline.py", "稳定的端到端 Python API"),
        ("config/pitch_mapping.csv", "20 种氨基酸音高及事实/推断状态"),
        ("tests/test_pipeline.py", "FASTA、PDB、CSV、十二音列回归测试"),
        (".streamlit/config.toml", "高对比主题、headless 与关闭统计提示"),
    ], [2800, 6560], size=9.2)
    doc.add_heading("13.1 Python API", level=2)
    add_code(doc, "from biomusic.pipeline import SonificationSettings, run_pipeline\nfrom biomusic.codec import decode_artifact\n\nsettings = SonificationSettings(\n    pitch_mode='可逆十二音列编解码', row_form='RI',\n    texture_density=6, counterpoint_strength=0.70\n)\nresult = run_pipeline('sample.fasta', fasta_bytes, settings)\nsequence, metadata, rows = decode_artifact('sample.mid', result.midi)\nassert sequence == ''.join(result.record.symbols)")

    doc.add_heading("14. 测试、验收与后续扩展", level=1)
    doc.add_heading("14.1 当前自动化测试", level=2)
    add_bullet(doc, "FASTA：完整生成、WAV RIFF、多轨 MIDI、多 Part MusicXML 与 GVR JSON。")
    add_bullet(doc, "十二音列：DNA/蛋白质、P/I/R/RI、康托排名边界、H_row/H_permutation/H_codec_domain。")
    add_bullet(doc, "逆向解码：GVR JSON、MusicXML、MIDI 均恢复原字符串；非法域严格拒绝。")
    add_bullet(doc, "PDB：空间声像、接触度与 NMA 可用性。")
    add_bullet(doc, "CSV：线粒体比例、QC 低通和转录组类型。")
    add_bullet(doc, "配器：所有事件必须属于九种古典乐器白名单。")
    add_bullet(doc, "多声部：至少 5 个 Part、至少 6 个 MIDI Track（含速度轨）、跨声部存在同时发声、同声部无意外重叠。")
    add_bullet(doc, "界面：示例与文本粘贴生成、六项指标、当前编制表、GVR 通过。")
    add_bullet(doc, "启动：全新临时用户环境不出现邮箱或 onboarding 提示。")
    doc.add_heading("14.2 扩展建议", level=2)
    add_para(doc, "可以在不改变验证器的前提下替换 Generate：接入本地 Transformer/LSTM/LLM，让模型只提出节奏、织体和主题发展；expected_pc、source_index、行游标与最终门控仍由确定性代码控制。专业声音方向可接入 FluidSynth/SF2、Muse Sounds、Kontakt 或 OSC/Max/MSP/SuperCollider；空间音频可升级为经受试者校准的 HRTF。")

    doc.add_heading("15. 术语表", level=1)
    add_table(doc, ["术语", "通俗解释"], [
        ("声学化 / Sonification", "把数据的变化系统地映射为可听参数"),
        ("音级 / Pitch class", "忽略八度后的 12 类音高，MIDI mod 12"),
        ("寄存器 / Register", "音符所在的高低音区"),
        ("声像 / Panning", "声音在左右声道之间的位置"),
        ("HVG", "高变特征/高变基因，信息变化较大的列"),
        ("NMA", "简正模式分析，用线性化模型描述协同振动模式"),
        ("GVR", "Generate–Verify–Repair，生成—验证—修复"),
        ("发布门控", "只有最终硬规则全部通过才提供结果"),
        ("Trace", "把音符连接回原始数据和映射规则的证书"),
        ("P/I/R/RI", "原形、倒影、逆行、逆行倒影十二音列形式"),
    ], [2500, 6860])

    doc.add_heading("16. 综述方法矩阵与本平台实现状态", level=1)
    doc.add_heading("16.1 综述提供的五条设计原则", level=2)
    reset_numbers()
    add_number(doc, "模态特异：DNA、表观、表达、蛋白、代谢与关联数据具有不同结构，不能共享一个万能字母—音符字典。")
    add_number(doc, "层级对应：最重要的生物组织轴控制最重要的音乐组织轴。序列顺序应控制时间，折叠或网络结构应控制织体、空间和形式。")
    add_number(doc, "感知节制：一次只选择少量可听辨变量；音高适合有序关系，节奏适合时间过程，音色适合类别，空间适合几何位置。")
    add_number(doc, "一致与可追溯：同批样本使用相同规则；报告数据、预处理、参数、软件、Trace 与验证结果。")
    add_number(doc, "比较与不确定性：优先在参考/突变、健康/疾病、处理/对照之间比较；低置信度不能被精美配器掩盖。")
    doc.add_heading("16.2 各数据模态如何进入当前六声部", level=2)
    add_table(doc, ["模态", "综述推荐的主要组织轴", "当前已实现", "仍待实现"], [
        ("DNA/RNA", "碱基、密码子、阅读框、起止密码子、重复主题", "序列→时间；碱基→音级；GC/嘌呤→音区与力度；六声部派生", "三阅读框独立声部、起止门控、k-mer 主题检测"),
        ("表观组", "坐标、峰高/宽、标记类别、比较状态", "列名识别；坐标→时间；信号→显著性与结构层", "多组蛋白标记独立声部、峰宽、状态比较"),
        ("转录组", "时间点、簇、通路、细胞状态、伪时间", "表达矩阵 QC；HVG/线粒体比例；多声部织体代理", "正规归一化、聚类、PCA、GO 背景层、伪时间分支"),
        ("蛋白/PDB", "序列、理化、二级结构、三维接触、振动", "氨基酸音高；疏水/电荷；HELIX/SHEET 时值；接触度；声像；粗粒化 NMA", "实验 NMR、真实力常数、构象系综、结合位点/修饰"),
        ("质谱/代谢", "峰、丰度、置信度、通量、通路、网络拓扑", "m/z、强度、丰度、置信度入口；竖琴与区块编配", "同位素通量、反应方向、通路图、中心性/模块"),
        ("GWAS/EWAS", "坐标、显著性、效应、LD、可信集、多基因结构", "坐标→时间；p 值→显著性；效应→音区；显著峰重音", "LD 区块和声场、可信集、跨性状/组织富集"),
    ], [1350, 2700, 3030, 2280], size=8.3)
    doc.add_heading("16.3 为什么采用“主旋律 + 派生声部”", level=2)
    add_para(doc, "综述反复指出，高维数据不应被压成孤立单音，但把全部变量同时映射又会形成不可理解的噪声。当前版本采取折中：原始数据点只生成一条信息最密集的前景主旋律，其余乐器从主事件、分组窗口、结构边界与显著峰派生。这样既产生真正复调，又能让每个派生音回到唯一的 parent_event_id。")
    add_callout(doc, "艺术层与分析层必须分开", "三度和声、主属低音、轮唱与厅堂反射用于可听性和音乐组织；它们不是由分子力学唯一决定的生物事实。Trace 中的 mapping_rule 会明确标注“派生”“区块”或“结构重音”，避免把配器选择误当成实验测量。", fill=PALE_GOLD)
    doc.add_heading("16.4 NMA、十二音列与 AI 的准确定位", level=2)
    add_para(doc, "PDB 粗粒化 NMA 提供相对模式背景，十二音列 GVR 提供可审计的音级硬约束，两者解决不同问题。前者近似结构动力学，后者约束符号作曲。当前版本不调用在线 LLM；未来可以让本地模型提出节奏、织体或主题发展，但 expected_pc、source_index、voice_id、parent_event_id 和发布门控仍应由确定性代码掌握。")
    doc.add_heading("16.5 综述来源与引用方式", level=2)
    add_para(doc, "本章的方法框架直接依据本地综述《Transform from genes to music》，特别是 genomic、epigenomic、transcriptomic、proteomic、metabolomic、association-study 与 cross-modality design principles 各节。正式论文中应引用综述所列的原始文献，而不是只引用本软件手册。")

    doc.add_page_break()
    doc.add_heading("附录 A：氨基酸音高映射", level=1)
    add_para(doc, "MIDI 与科学音高来自 config/pitch_mapping.csv。explicit 表示补充材料明确给出；explicit_from_TS 表示来自 TS 连接段；inferred 表示根据复现逻辑推断；extended_inference 是为通用 20 氨基酸输入补充的软件映射。")
    pitch_rows = [
        ("A", "Ala", "81 / A5", "inferred"), ("R", "Arg", "74 / D5", "explicit"),
        ("N", "Asn", "65 / F4", "extended_inference"), ("D", "Asp", "62 / D4", "explicit"),
        ("C", "Cys", "61 / C#4", "extended_inference"), ("Q", "Gln", "63 / Eb4", "explicit"),
        ("E", "Glu", "76 / E5", "explicit"), ("G", "Gly", "79 / G5", "explicit"),
        ("H", "His", "71 / B4", "explicit"), ("I", "Ile", "68 / Ab4", "extended_inference"),
        ("L", "Leu", "80 / Ab5", "explicit"), ("K", "Lys", "78 / F#5", "explicit"),
        ("M", "Met", "64 / E4", "explicit"), ("F", "Phe", "77 / F5", "explicit"),
        ("P", "Pro", "75 / Eb5", "explicit"), ("S", "Ser", "70 / Bb4", "explicit"),
        ("T", "Thr", "72 / C5", "explicit_from_TS"), ("W", "Trp", "68 / Ab4", "extended_inference"),
        ("Y", "Tyr", "90 / F#6", "explicit"), ("V", "Val", "73 / Db5", "explicit"),
        ("X", "Unknown", "60 / C4", "fallback"),
    ]
    add_table(doc, ["符号", "名称", "MIDI / 音高", "状态"], pitch_rows, [1150, 2200, 2600, 3410], size=9.2)

    doc.add_heading("附录 B：Trace CSV 字段", level=1)
    add_table(doc, ["字段", "含义"], [
        ("event_id", "最终事件的全局唯一序号"), ("voice_id/role", "独立声部与音乐功能"),
        ("parent_event_id", "派生声部指向主旋律事件；主旋律为空"), ("source_index/source_label", "原始位置与人类可读标签"),
        ("symbol", "氨基酸、碱基、cell 或 peak"), ("onset_quarter/duration_quarter", "以四分音符为单位的起拍与时值"),
        ("midi/pitch_class", "实际音高与实际音级"), ("expected_pc", "生物映射或十二音列证书要求的音级"),
        ("velocity/pan", "力度和左右声像"), ("timbre", "九种古典配器内部名称"),
        ("row_position/row_form", "十二音列位置与呈现形式；非序列模式为空"), ("codec_block/is_codec_carrier", "载体块号与是否参与解码"),
        ("status", "retained 或 repaired"),
        ("mapping_rule", "该事件的音高、派生或结构规则摘要"), ("hydropathy/charge/contact/value/uncertainty", "进入音乐映射的主要数值特征"),
    ], [3300, 6060], size=9.2)

    doc.add_heading("附录 C：GVR JSON 结构", level=1)
    add_code(doc, "{\n  'metadata': {'record_name': '...', 'audio': {...}, 'nma': {...}},\n  'gvr': {\n    'passed': true,\n    'checks': {'H_row': true, 'H_permutation': true, 'H_codec_domain': true},\n    'tone_rows': [[0, ...], ...],\n    'codec': {\n      'codec_version': 'biosound-cantor-v1', 'data_type': 'protein',\n      'alphabet': 'ACDEFGHIKLMNPQRSTVWY*', 'block_size': 6,\n      'row_form': 'RI', 'original_length': 83, 'pad_length': 1\n    }\n  }\n}")

    doc.add_heading("附录 D：科研发布前检查清单", level=1)
    for item in [
        "原始输入文件与校验值已归档。", "软件版本、配置文件与随机种子已记录。",
        "抽样步长、速度、拍号、音域、调式与音高模式已报告。", "PDB 接触/NMA 参数和科学边界已写入方法。",
        "CSV 的矩阵方向、MT 基因命名和 HVG 方法已经独立复核。", "Trace CSV 和 GVR JSON 与音频/谱面一起发布。",
        "没有把网页预览音色描述成真实乐团采样。", "没有把声学化结果描述成诊断或结构证明。",
        "盲听评价报告了样本量、问题、随机化和评分标准。", "最终 MusicXML 已在 MuseScore 中人工检查拍号、音符与配器文字。",
    ]:
        add_bullet(doc, "□ " + item)

    doc.add_page_break()
    doc.add_heading("附录 E：论文灵感与本平台实现的对应", level=1)
    add_para(doc, "方法参考：Dai et al., Verifier-Guided Twelve-Tone Composition: A Generate–Verify–Repair Harness for Symbolic Music Generation, arXiv:2607.11334v2, 2026；本地综述《Transform from genes to music》；以及本地 Spinning Melodies 复现中的 pitch_mapping.csv 逻辑。", italic=True)
    add_table(doc, ["论文/方法思想", "平台实现", "差异与边界"], [
        ("规范化事件记录", "MusicEvent + voice/role/parent + Trace CSV", "实现六声部并保留派生关系"),
        ("Generate", "可复现规则生成器", "当前不调用在线 LLM"),
        ("Deterministic Verify", "H_mapping 等最终事件复检", "只声明已实现规则"),
        ("Row-preserving Repair", "按声部换八度保音级、修时值、顺延内部重叠", "没有 LLM patch/replan"),
        ("Final release gate", "violations_after 为空才发布", "失败会明确报错"),
        ("Trace", "每个音符保存来源、证书和状态", "计划层历史仍有限"),
        ("十二音列 P/I/R/RI", "分块进制数 + 康托编解码 + 可逆呈现形式", "V1 承载数据；生物特征控制其他维度"),
    ], [2600, 3300, 3460], size=9.0)

    # Keep core properties explicit for deterministic updates in Word.
    props = doc.core_properties
    props.title = "BioSound GVR 多声部管弦乐版平台功能与原理说明"
    props.subject = "综述驱动的模态特异生物数据声学化、六声部古典配器与 GVR"
    props.author = "BioSound GVR Project"
    props.keywords = "sonification, GVR, multi-omics, FASTA, PDB, NMA, multi-track MIDI, multi-part MusicXML, classical orchestration"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
