from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"C:\Users\34591\Desktop\音乐基因转化\复现材料\复现评价.docx")
NAVY = "0B2545"; BLUE = "2E74B5"; DARK = "1F4D78"; LIGHT = "E8EEF5"; PALE = "F4F6F9"; GOLD = "7A5A00"; RED = "9B1C1C"

def set_font(run, size=None, bold=None, color=None, east="Microsoft YaHei"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def margins(cell, top=80, start=120, bottom=80, end=120):
    tcPr = cell._tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for side, val in [("top",top),("start",start),("bottom",bottom),("end",end)]:
        el = tcMar.find(qn("w:"+side))
        if el is None: el = OxmlElement("w:"+side); tcMar.append(el)
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")

def width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.find(qn("w:tcW"))
    if tcW is None: tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
    tcW.set(qn("w:w"), str(dxa)); tcW.set(qn("w:type"), "dxa")

def make_table(doc, headers, rows, widths, header_fill=LIGHT):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.autofit = False
    tblPr = t._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW")); tblW.set(qn("w:w"), "9360"); tblW.set(qn("w:type"), "dxa")
    ind = OxmlElement("w:tblInd"); ind.set(qn("w:w"), "120"); ind.set(qn("w:type"), "dxa"); tblPr.append(ind)
    grid = t._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; width(c,widths[i]); margins(c); shade(c,header_fill); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: set_font(r,9.5,True,NAVY)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=str(val); width(cells[i],widths[i]); margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.08
                for r in p.runs: set_font(r,9.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t

def callout(doc, label, text, fill=PALE, color=NAVY):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.autofit=False
    c=t.cell(0,0); width(c,9360); margins(c,140,180,140,180); shade(c,fill)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.15
    r=p.add_run(label+"  "); set_font(r,10.5,True,color)
    r=p.add_run(text); set_font(r,10.5,False,NAVY)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def labeled(doc, label, text):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5); p.paragraph_format.line_spacing=1.18
    r=p.add_run(label+"："); set_font(r,10.5,True,DARK)
    r=p.add_run(text); set_font(r,10.5)

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.75); sec.left_margin=Inches(0.85); sec.right_margin=Inches(0.85); sec.header_distance=Inches(0.35); sec.footer_distance=Inches(0.35)
styles=doc.styles
normal=styles["Normal"]; normal.font.name="Calibri"; normal._element.rPr.rFonts.set(qn("w:eastAsia"),"Microsoft YaHei"); normal.font.size=Pt(10.5); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.18
for name,size,color,before,after in [("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,14,7),("Heading 3",12,DARK,10,5)]:
    s=styles[name]; s.font.name="Calibri"; s._element.rPr.rFonts.set(qn("w:eastAsia"),"Microsoft YaHei"); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
header=sec.header.paragraphs[0]; header.text="Spinning Melodies｜复现评价"; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs: set_font(r,8.5,False,"6B7280")
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=footer.add_run("面向非生物专业的音乐入门读者  ·  "); set_font(r,8,False,"6B7280")
field=OxmlElement("w:fldSimple"); field.set(qn("w:instr"),"PAGE"); footer._p.append(field)

# Editorial cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(58); p.paragraph_format.space_after=Pt(10)
r=p.add_run("复现评价"); set_font(r,30,True,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8)
r=p.add_run("从蛋白序列到长笛旋律：当前 MusicXML 与原始录音的对比"); set_font(r,14,False,DARK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(34)
r=p.add_run("给音乐初学者的听觉导读与改进建议"); set_font(r,11,False,"6B7280")
callout(doc,"一句话结论", "当前成果已经成功复现论文的音高映射、构建体身份和曲式骨架，适合作为教学演示；但它尚不是原录音的逐音复刻，尤其在时长、节奏变化、呼吸、力度和音区塑形方面仍有明显距离。", "EAF2F8", NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(24)
r=p.add_run("评价日期：2026年7月15日"); set_font(r,9.5,False,"6B7280")
doc.add_page_break()

doc.add_heading("1. 这次究竟复现了什么？",1)
doc.add_paragraph("可以把这项工作理解成：把一串蛋白质字母翻译成一支长笛能演奏的旋律。字母先变成固定音高，再组合成 H、A、B 和 TS 四类乐句。当前 MusicXML 主要复现的是这套“翻译语法”，而原始 MP3 还包含演奏者与作曲者加入的速度伸缩、呼吸、力度、音区和乐句塑形。")
labeled(doc,"H", "四首共同的48音开头，像一段相同的序言。")
labeled(doc,"A", "以丙氨酸（Ala）为主的紧凑主题；A音高频反复，听起来较集中、执拗。")
labeled(doc,"B", "更长、音高种类更多的主题；相较 A 更容易形成流动、抒情的线条。")
labeled(doc,"TS", "结束或连接乐句的 C–B♭ 长音；每个音持续6个四分音符拍。")
labeled(doc,"CODA", "第4首末尾的 E–B–B–B，像刚结束又准备回到第一首开头。")

doc.add_heading("2. 评价方法与边界",1)
doc.add_paragraph("本评价使用4个原始 MP3、4个当前 MusicXML，以及补充材料第14–15页给出的映射规则。MP3 时长由本地媒体信息读取；MusicXML 的速度、拍号、音符时值和音域由文件结构直接统计。由于论文没有公开完整总谱，本评价不把“与录音逐音一致”作为当前版本已经达到的事实。")
callout(doc,"重要边界", "以下评分用于判断教学复现的完成度，不是统计学实验结果。没有完整原谱时，无法仅凭录音断言每一个音符究竟是作曲设定还是演奏处理。", "FFF8E8", GOLD)

doc.add_heading("3. 总体评价",1)
make_table(doc,["评价维度","当前水平","解释"],[
    ("规则与构建体对应","较好（8/10）","编号、HAB₃/HA₃B、拍号、速度、TS 和 coda 均与补充说明一致。"),
    ("旋律骨架","良好（7/10）","字母到音高的映射清楚，H/A/B 主题可在谱面中辨认。"),
    ("原录音听觉接近度","有限（3/10）","主要差距来自时长、节奏、呼吸、力度、音区和演奏表情。"),
    ("初学者教学价值","较好（8/10）","字母写在音符下方，适合边看谱边听主题如何形成。"),
    ("综合定位","教学版重建（6/10）","可以展示方法，但不宜称为原作完整复刻。"),
],[2100,1800,5460])

doc.add_heading("4. 最直观的差距：曲目长度",1)
doc.add_paragraph("原录音明显比当前乐谱长。原因不是速度写错，而是当前版本把大多数氨基酸都处理成连续八分音符，缺少原作中的长短变化、停顿、呼吸、重复和速度伸缩。")
make_table(doc,["曲目","构建体","原录音","当前谱面","时长覆盖率","判断"],[
    ("01","HAB₃","58.0秒","37.9秒","65%","四首中最接近；仍偏短约20秒。"),
    ("02","HA₃B","76.6秒","36.7秒","48%","主题正确，但乐句被明显压缩。"),
    ("03","HAB₃","128.5秒","55.9秒","43%","虽为加长版，展开程度仍不到原录音一半。"),
    ("04","HA₃B","203.6秒","53.3秒","26%","差距最大；原录音显然有更充分的发展与停顿。"),
],[800,1050,1200,1200,1250,3860])
p=doc.add_paragraph("说明：当前谱面时长按 MusicXML 中写入的速度和小节时值计算；覆盖率=当前谱面时长÷原录音时长。录音首尾静音和演奏速度伸缩可能带来小幅误差，但不会改变总体结论。")
p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(8)
for r in p.runs: set_font(r,9,False,"6B7280")

doc.add_heading("5. 分曲评价",1)
doc.add_heading("01｜HAB₃｜Vivo，6/8｜形成纤维",2)
labeled(doc,"做得好", "共同 H 前缀、A 后接三个 B 主题、Vivo 速度和末尾 TS 都很清楚。B 主题占比高，容易听成较连续的旋律线。")
labeled(doc,"主要差距", "当前版本节奏几乎等长，缺少原录音中的呼吸和动态弧线，因此“流动但有形状”的感觉还不够充分。")
labeled(doc,"当前判断", "最适合用作第一轮教学示范，也是四首中结构与时长最接近原录音的一首。")

doc.add_heading("02｜HA₃B｜Andantino，3/4｜不形成纤维",2)
labeled(doc,"做得好", "三个 A 主题使 A 音及邻近音反复出现，较容易形成紧凑、局域、略带执拗的听感；这与论文对 A-rich 材料局部聚集的类比方向一致。")
labeled(doc,"主要差距", "当前谱面只有原录音约48%的长度。若没有力度对比、重音和呼吸位置，A 的‘攻击性’容易退化成单纯重复。")
labeled(doc,"当前判断", "物理概念表达是成立的，但音乐性仍需通过节奏和重音来加强。")

doc.add_heading("03｜HAB₃｜Vivo，6/8｜形成纤维（加长版）",2)
labeled(doc,"做得好", "主题数量确实增加，HAB₃ 的 B-rich 特征被保留，曲式方向正确。")
labeled(doc,"主要差距", "原录音约129秒，当前仅约56秒。加长主要靠增加主题数量，而不是通过节奏变形、音区迁移、力度起伏和呼吸来发展主题。")
labeled(doc,"当前判断", "更像‘扩充后的材料清单’，尚未达到原录音那种完整的第二旋律发展。")

doc.add_heading("04｜HA₃B｜Andantino，3/4｜不形成纤维（加长版）",2)
labeled(doc,"做得好", "A-rich 结构明确，结尾先出现 TS，再以 E–B–B–B 的 coda 指回第一首开头；这一设计非常适合初学者理解‘未真正结束’的悬念。")
labeled(doc,"主要差距", "原录音约204秒，当前约53秒，只有约四分之一。补充材料特别指出第4首后段音区处理更自由，而当前各曲基本共用固定音区。")
labeled(doc,"当前判断", "结尾概念完成度高，但主体发展不足，是下一轮最应优先细化的一首。")

doc.add_page_break()
doc.add_heading("6. 为什么现在听起来更像‘编码’，还不像完整作品？",1)
make_table(doc,["音乐要素","当前实现","原录音中可期待的作用","改进方向"],[
    ("节奏","多数为八分音符","制造停顿、推动与句尾感","按录音逐句转录长短时值。"),
    ("呼吸","未明确标注","决定长笛乐句在哪里分开","标出换气点，并留出真实休止。"),
    ("力度","基本固定","让 A 更紧张、B 更有流动弧线","加入 p–f、渐强和渐弱。"),
    ("音区","四首映射较固定","同一主题在高低音区会产生不同张力","依据录音记录每次主题移高或移低。"),
    ("奏法","缺少重音/连奏","区分攻击性与抒情性","A 增加重音和短连线；B 使用更长连奏。"),
    ("人声化处理","MIDI较机械","真实长笛有起音、气息和速度伸缩","用更好的长笛音源或真人演奏。"),
],[1250,1900,3000,3210])

doc.add_heading("7. 给音乐初学者的对比聆听方法",1)
for title,text in [
    ("第一遍：只听开头", "四首开头是否都像同一句话？这就是共同的 H 前缀。"),
    ("第二遍：寻找 A", "听到 A 音大量反复、旋律集中在较窄区域时，把它标成 A 主题。"),
    ("第三遍：寻找 B", "比较 B 是否更长、更流动、音高轮廓是否更有变化。"),
    ("第四遍：听结尾", "01–03 应由 C–B♭ 长音收束；04 在长音后还有 E–B–B–B。"),
    ("第五遍：比较呼吸", "原录音在哪里停顿、渐强或换气，而 MIDI 没有？这些位置就是下一轮需要转录的音乐信息。"),
]:
    doc.add_heading(title,3); doc.add_paragraph(text)

doc.add_heading("8. 下一轮改进建议（按优先级）",1)
make_table(doc,["优先级","任务","完成标准"],[
    ("1","逐首标记原录音的主题起止时间","H、A、B、TS、coda 均有时间戳；段落顺序可复核。"),
    ("2","转录真实节奏与休止","总时长与原录音差距控制在±5%，主要句尾与换气一致。"),
    ("3","补入力度、连奏、重音和呼吸","谱面能解释为什么 A 较紧张、B 较抒情。"),
    ("4","校准音区变化","同一主题每次出现的八度位置与录音一致。"),
    ("5","优化播放音色","使用更自然的长笛音源，或邀请演奏者录制。"),
    ("6","做盲听小测试","让初学者在不知道标签时判断哪首更紧凑、哪首更流动。"),
],[900,3300,5160])

doc.add_heading("9. 最终结论",1)
callout(doc,"适合怎样使用", "当前版本适合课堂演示、规则验证和继续编辑：读者能看到蛋白字母如何变成音符，也能分辨 A-rich 与 B-rich 的基本差异。", "EAF2F8", NAVY)
callout(doc,"暂时不宜怎样表述", "不宜称为对四段原始录音的完整或逐音复刻。更准确的名称是：基于补充规则的可复跑 MusicXML 教学重建。", "FFF1F1", RED)
doc.add_paragraph("若下一轮优先完成原录音的主题时间戳、节奏与呼吸转录，整体听觉接近度会得到最大提升；其中第4首应作为首要修订对象。")

doc.add_heading("附：本次使用的本地材料",1)
for text in [
    "补充说明：NIHMS424798-supplement-05.pdf（第14–15页）",
    "原录音：NIHMS424798-supplement-01.mp3 至 04.mp3",
    "复现谱：01_HAB3、02_HA3B、03_HAB3、04_HA3B 的 rule_reconstruction.musicxml",
    "评价方式：录音元数据时长 + MusicXML 结构统计 + 补充规则一致性检查。",
]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(text)

doc.core_properties.title="复现评价"
doc.core_properties.subject="Spinning Melodies 原始音频与 MusicXML 规则复现的对比评价"
doc.core_properties.author="Codex"
doc.save(OUT)
print(OUT)
